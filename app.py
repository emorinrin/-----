# main.py
import streamlit as st
import google.generativeai as genai
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
import time
import io
import docx # docxファイルを扱うためにインポート

# --- アプリケーションの基本設定 ---
st.set_page_config(
    page_title="長文PDF翻訳アプリ",
    page_icon="📖",
    layout="wide"
)

# --- 関数定義 ---

@st.cache_data(show_spinner=False)
def extract_text_from_pdf(uploaded_file):
    """
    アップロードされたPDFファイルからテキストを抽出する関数
    """
    try:
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        return text
    except Exception as e:
        st.error(f"PDFの読み込み中にエラーが発生しました: {e}")
        return None

def translate_text(text_chunk, api_key, model_name="gemini-1.5-flash"):
    """
    Google Gemini APIを使用してテキストチャンクを翻訳する関数
    """
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        # プロンプトの設計：より自然な翻訳を促す指示を追加
        prompt = f"""
        あなたはプロの翻訳家です。以下の英文を、文脈を考慮し、自然で流暢な日本語に翻訳してください。
        専門用語は正確に訳し、元の文章の意図やニュアンスを忠実に再現してください。
        原文に関する解説などの追加情報や前置きは不要です。

        ---
        English Text:
        {text_chunk}
        ---

        Japanese Translation:
        """
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        # APIエラーの詳細をユーザーにフィードバック
        return f"翻訳中にエラーが発生しました: {str(e)}"

# --- Streamlit UI ---

st.title("📖 長文PDF翻訳アプリ")
st.markdown("100ページを超えるような長文の英文PDFファイルを、GoogleのGeminiモデルを使って日本語に翻訳します。")
st.markdown("---")

# --- サイドバー ---
with st.sidebar:
    st.header("設定")
    api_key = st.text_input("Google AI Studio APIキー", type="password")
    st.markdown("[APIキーの取得方法](https://aistudio.google.com/app/apikey)", unsafe_allow_html=True)
    
    model_name = st.selectbox(
        "使用するモデル",
        ("gemini-1.5-flash", "gemini-1.5-pro", "gemini-2.5-pro"),
        help="通常は`gemini-1.5-flash`が高速かつ十分な性能でおすすめです。"
    )
    
    chunk_size = st.slider(
        "チャンクサイズ", 500, 5000, 1500, 100,
        help="一度に翻訳する文字数。大きいほど文脈を維持できますが、APIの制限に注意してください。"
    )
    chunk_overlap = st.slider(
        "オーバーラップ", 0, 500, 200, 50,
        help="チャンク間の重複文字数。文脈の連続性を高めます。"
    )

# --- メインコンテンツ ---
uploaded_file = st.file_uploader(
    "翻訳したい英文のPDFファイルをアップロードしてください",
    type="pdf",
    accept_multiple_files=False
)

if 'translated_text' not in st.session_state:
    st.session_state.translated_text = ""
if 'original_text' not in st.session_state:
    st.session_state.original_text = ""

if uploaded_file:
    if st.button("翻訳開始", type="primary"):
        # 初期化
        st.session_state.translated_text = ""
        st.session_state.original_text = ""

        # 1. PDFからテキスト抽出
        with st.spinner("PDFからテキストを抽出中..."):
            original_text = extract_text_from_pdf(uploaded_file)
            st.session_state.original_text = original_text

        if original_text:
            st.success("テキストの抽出が完了しました。翻訳を開始します。")

            # 2. テキストをチャンクに分割
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
            )
            chunks = text_splitter.split_text(original_text)
            total_chunks = len(chunks)

            # 3. チャンクを順番に翻訳
            progress_bar = st.progress(0, "翻訳の進捗")
            translated_chunks = []
            
            # リアルタイム表示用のプレースホルダー
            realtime_display = st.empty()
            
            start_time = time.time()

            for i, chunk in enumerate(chunks):
                # 翻訳実行
                translated_chunk = translate_text(chunk, api_key, model_name)
                translated_chunks.append(translated_chunk)
                
                # 進捗とリアルタイム表示の更新
                progress_value = (i + 1) / total_chunks
                progress_text = f"翻訳の進捗: {i+1}/{total_chunks} チャンク完了"
                progress_bar.progress(progress_value, text=progress_text)
                
                # 翻訳済みテキストを結合してリアルタイム表示
                realtime_display.text_area(
                    "翻訳結果（リアルタイム更新）", 
                    "".join(translated_chunks), 
                    height=400
                )
                
                # APIのレートリミット対策
                time.sleep(1) 

            end_time = time.time()
            processing_time = end_time - start_time
            
            st.session_state.translated_text = "".join(translated_chunks)
            progress_bar.empty() # 完了したらプログレスバーを消す
            st.success(f"翻訳が完了しました！（処理時間: {processing_time:.2f}秒）")

# 翻訳結果の表示とダウンロード
if st.session_state.translated_text:
    st.subheader("最終翻訳結果")
    st.text_area(
        "翻訳全文",
        st.session_state.translated_text,
        height=600
    )
    
    st.subheader("ダウンロード")

    # --- docxファイル生成ロジック ---
    doc = docx.Document()
    doc.add_heading(f"{uploaded_file.name} の翻訳結果", 0)
    doc.add_paragraph(st.session_state.translated_text)
    
    # メモリ上でdocxファイルを扱う
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    # --------------------------------

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="テキストファイル (.txt)",
            data=st.session_state.translated_text.encode('utf-8'),
            file_name=f"{uploaded_file.name.replace('.pdf', '')}_translated.txt",
            mime="text/plain",
            use_container_width=True
        )
    with col2:
        st.download_button(
            label="Wordドキュメント (.docx)",
            data=bio,
            file_name=f"{uploaded_file.name.replace('.pdf', '')}_translated.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


    with st.expander("原文を表示"):
        st.text_area("抽出された原文", st.session_state.original_text, height=400)
